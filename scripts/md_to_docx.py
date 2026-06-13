"""
Convierte el informe unificado .md a .docx con formato específico:
- Times New Roman 12
- Interlineado 1.15
- Justificado
- Jerarquía de títulos respetada (H1 -> Heading 1, H2 -> Heading 2, etc.)
- Sin portada, sin tabla de contenidos, sin resumen, sin referencias

Uso: python scripts/md_to_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Configuración ---
INPUT_MD = Path(r"d:\tareas antiguas\CdeDatos\proyecto-autos\reports\06_INFORME_UNIFICADO.md")
OUTPUT_DIR = Path(r"d:\tareas antiguas\CdeDatos\proyecto-autos\entregas\HitoFinal")
OUTPUT_DOCX = OUTPUT_DIR / "06_INFORME_UNIFICADO.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.15


def set_paragraph_format(paragraph, font_name=FONT_NAME, font_size=FONT_SIZE,
                         bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Aplica formato base a un párrafo."""
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        run.font.italic = italic
        # Para que Word respete la fuente en todos los idiomas
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)


def set_heading_format(paragraph, level):
    """Aplica formato a un heading."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(18) if level <= 2 else Pt(12)
    pf.space_after = Pt(6)

    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(13), 4: Pt(12)}
    size = sizes.get(level, Pt(12))

    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = size
        run.font.bold = True
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)


def add_table_to_doc(doc, header_row, data_rows):
    """Agrega una tabla con formato al documento."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'

    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= cols:
                break
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Espacio después de la tabla
    doc.add_paragraph()


def clean_inline_markdown(text):
    """Limpia formato inline de markdown para texto plano."""
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Code inline
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Links [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Images
    text = re.sub(r'!\[.*?\]\(.+?\)', '[Imagen]', text)
    return text


def add_rich_paragraph(doc, text, is_list=False, list_marker=""):
    """Agrega un párrafo con formato rich (negritas, cursivas, código)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)

    if is_list:
        pf.left_indent = Cm(1)
        pf.first_line_indent = Cm(-0.5)
        text = list_marker + " " + text

    # Parse inline formatting
    # Split por bold, italic, code
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        else:
            # Handle links and images
            subparts = re.split(r'(!\[.*?\]\(.*?\)|\[.*?\]\(.*?\))', part)
            for subpart in subparts:
                if subpart.startswith('!['):
                    run = p.add_run('[Imagen - ver archivo .md]')
                    run.font.italic = True
                elif subpart.startswith('['):
                    link_text = re.match(r'\[(.+?)\]', subpart)
                    if link_text:
                        run = p.add_run(link_text.group(1))
                    else:
                        run = p.add_run(subpart)
                else:
                    run = p.add_run(subpart)

            continue

        # Set font for formatted runs
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if not (part.startswith('`') and part.endswith('`')):
            rFonts.set(qn('w:ascii'), FONT_NAME)
            rFonts.set(qn('w:hAnsi'), FONT_NAME)

    # Set default font for all runs without explicit font
    for run in p.runs:
        if run.font.name is None:
            run.font.name = FONT_NAME
        if run.font.size is None:
            run.font.size = FONT_SIZE
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if not rFonts.get(qn('w:ascii')):
            rFonts.set(qn('w:ascii'), FONT_NAME)
            rFonts.set(qn('w:hAnsi'), FONT_NAME)

    return p


def convert_md_to_docx():
    """Función principal de conversión."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Configurar estilo normal del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- Parsear el markdown ---
    i = 0
    skip_section = False
    in_code_block = False
    code_lines = []
    in_table = False
    table_header = []
    table_rows = []

    # Secciones a excluir
    skip_titles = [
        "Tabla de Contenidos",
        "Resumen",
        "13. Referencias"
    ]

    # También saltar el encabezado YAML/metadata al inicio
    skip_header = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Saltar encabezado del documento (título, autor, etc. antes del primer ---)
        if skip_header:
            if stripped == '---' and i > 0:
                skip_header = False
                i += 1
                continue
            elif stripped.startswith('# '):
                skip_header = False
                # No incrementar i, procesar este heading
            else:
                i += 1
                continue

        # Manejo de bloques de código
        if stripped.startswith('```'):
            if in_code_block:
                # Cerrar bloque de código
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # Manejo de tablas
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                # Inicio de tabla
                in_table = True
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_header = cells
                i += 1
                continue
            elif re.match(r'^[\|\s\-:]+$', stripped):
                # Línea separadora de tabla, saltar
                i += 1
                continue
            else:
                # Fila de datos
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
                continue
        elif in_table:
            # Fin de tabla
            if table_header:
                add_table_to_doc(doc, table_header, table_rows)
            in_table = False
            table_header = []
            table_rows = []
            # No incrementar, procesar la línea actual

        # Líneas vacías
        if not stripped:
            i += 1
            continue

        # Separadores ---
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)

            # Verificar si es una sección a excluir
            should_skip = False
            for skip_title in skip_titles:
                if skip_title.lower() in title.lower():
                    should_skip = True
                    break

            if should_skip:
                skip_section = True
                i += 1
                continue
            else:
                # Si encontramos un nuevo heading de nivel 1 o 2 después de skip, salir
                if skip_section and level <= 2:
                    skip_section = False
                elif skip_section:
                    i += 1
                    continue

            # Agregar heading
            clean_title = clean_inline_markdown(title)
            p = doc.add_heading(clean_title, level=level)
            set_heading_format(p, level)
            i += 1
            continue

        # Si estamos en una sección excluida, saltar
        if skip_section:
            i += 1
            continue

        # Blockquotes
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            if text:
                p = add_rich_paragraph(doc, text)
                p.paragraph_format.left_indent = Cm(1)
                for run in p.runs:
                    run.font.italic = True
            i += 1
            continue

        # Imágenes standalone
        if stripped.startswith('!['):
            p = doc.add_paragraph()
            run = p.add_run('[Imagen - ver archivo .md para la visualización]')
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = LINE_SPACING
            i += 1
            continue

        # Listas con viñetas
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            add_rich_paragraph(doc, text, is_list=True, list_marker="•")
            i += 1
            continue

        # Listas numeradas
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            number = num_match.group(1)
            text = num_match.group(2)
            add_rich_paragraph(doc, text, is_list=True, list_marker=f"{number}.")
            i += 1
            continue

        # Párrafo normal
        add_rich_paragraph(doc, stripped)
        i += 1

    # Si quedó una tabla pendiente al final
    if in_table and table_header:
        add_table_to_doc(doc, table_header, table_rows)

    # Guardar
    doc.save(str(OUTPUT_DOCX))
    print(f"Documento generado en: {OUTPUT_DOCX}")
    print(f"Tamaño: {OUTPUT_DOCX.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    convert_md_to_docx()
